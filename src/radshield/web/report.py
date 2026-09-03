"""Word (.docx) shielding report generation.

Turns computed results into the skeleton of the kind of report a physicist
actually delivers: auto-filled data tables where the app genuinely has the
numbers, and clearly-marked placeholders everywhere else (facility identity,
project-specific assumptions, floor-plan figures, narrative findings and
recommendations). This is deliberately not an attempt to write the narrative
sections for the reviewer -- the app has no basis for "prior use", "which
walls were inspected", or "why gaps were found", and guessing at them would
be worse than leaving an obvious blank.

Figures are a placeholder rather than an embedded, auto-annotated floor plan
image: the app's own canvas already shows source and point markers, wall
color and opacity are user-adjustable specifically to make a clean
screenshot (see the wall color/opacity controls), and reproducing that
rendering here in Python would duplicate the real renderer rather than
reuse it.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from ..engine.evaluate import PointResult
from ..model.project import Project


class ReportError(RuntimeError):
    """Raised when a report cannot be generated."""


def _docx():
    """Import python-docx lazily so the physics package stays dependency-free."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise ReportError(
            "python-docx is required for report generation. Install it with: "
            "pip install python-docx"
        ) from exc
    return docx


_PLACEHOLDER_RGB = (0xB0, 0x00, 0x00)

_GENERAL_ASSUMPTIONS = (
    "All holes in barriers for pipes, conduits, and louvers shall be provided with "
    "baffles so that the radiation transmitted through them does not exceed that "
    "transmitted through the surrounding barriers.",
    "Windows and doors shall offer the same degree of attenuation as that required "
    "of the barrier in which they are located.",
    "Joints between lead sheets should be constructed such that there is an overlap "
    "of at least 0.4 inches.",
    "Door and window frames are to offer the same attenuation as that required of "
    "the barrier in which they are located.",
    "Any double doors requiring shielding require a shielded door astragal.",
    "All lead shielding, unless otherwise specified, shall extend to a height of 7 ft.",
)

_METHOD_LABELS = {
    "tg108": "TG-108 (nuclear medicine)",
    "ncrp147": "NCRP 147 (x-ray / fluoroscopy)",
    "ncrp147_ct": "NCRP 147 (CT)",
}


def _placeholder(paragraph: Any, text: str) -> None:
    """Append a run styled as an obvious "fill this in" marker."""
    from docx.shared import RGBColor

    run = paragraph.add_run(f"[{text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(*_PLACEHOLDER_RGB)


def _placeholder_paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    _placeholder(p, text)


def _set_cell_text(cell: Any, text: str, *, bold: bool = False) -> None:
    from docx.shared import Pt

    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def _methods_in_use(project: Project) -> list[str]:
    seen = {src.method for src in project.sources}
    return [_METHOD_LABELS.get(m, m) for m in sorted(seen)]


def _goal_lines(project: Project) -> list[str]:
    """Design goal bullets, reflecting only the methodologies actually in use."""
    from ..physics.limits import ncrp147_goal, tg108_goal

    methods = {src.method for src in project.sources}
    lines: list[str] = []
    if "tg108" in methods:
        lines.append(
            f"TG-108: uncontrolled areas < {tg108_goal('uncontrolled').value:g} uSv/week "
            f"effective dose equivalent; controlled areas < {tg108_goal('controlled').value:g} uSv/week."
        )
    if {"ncrp147", "ncrp147_ct"} & methods:
        lines.append(
            f"NCRP 147: uncontrolled areas < {ncrp147_goal('uncontrolled').value:g} mGy/week "
            f"air kerma; controlled areas < {ncrp147_goal('controlled').value:g} mGy/week."
        )
    if not lines:
        lines.append(
            f"TG-108: uncontrolled < {tg108_goal('uncontrolled').value:g} uSv/week, "
            f"controlled < {tg108_goal('controlled').value:g} uSv/week. "
            f"NCRP 147: uncontrolled < {ncrp147_goal('uncontrolled').value:g} mGy/week, "
            f"controlled < {ncrp147_goal('controlled').value:g} mGy/week."
        )
    return lines


def _governing_method(result: PointResult):
    governing = next((m for m in result.methods if m.method == "combined"), None)
    if governing is None and result.methods:
        governing = result.methods[0]
    return governing


def _add_targets_table(doc: Any, results: list[PointResult], materials: list[str]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    headers = [
        "Target", "Floor", "Method", "P", "T", "P/T", "Unshielded DR", "Shielded DR", "% of Goal",
    ] + [f"{m.title()} required (mm)" for m in materials]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True)

    for result in results:
        governing = _governing_method(result)
        row = table.add_row().cells
        _set_cell_text(row[0], result.label)
        _set_cell_text(row[1], result.floor_name)
        if governing is None:
            _set_cell_text(row[2], "unavailable")
            for cell in row[3:]:
                _set_cell_text(cell, "—")
            continue
        p_over_t = (
            f"{governing.goal_value / result.occupancy:.4g}" if result.occupancy > 0 else ""
        )
        unit = "uSv/wk" if "uSv" in governing.quantity else "mGy/wk"
        pct = (
            "" if governing.required_transmission == float("inf")
            else f"{100.0 / governing.required_transmission:.1f}%"
        )
        _set_cell_text(row[2], governing.method)
        _set_cell_text(row[3], f"{governing.goal_value:g} {unit}")
        _set_cell_text(row[4], f"{result.occupancy:g}")
        _set_cell_text(row[5], f"{p_over_t} {unit}" if p_over_t else "")
        _set_cell_text(row[6], f"{governing.unshielded_total:.4g} {unit}")
        _set_cell_text(row[7], f"{governing.total:.4g} {unit}")
        _set_cell_text(row[8], pct)
        for i, material in enumerate(materials):
            required = result.governing_thickness_mm.get(material)
            text = "unavailable" if governing.unavailable.get(material) else (
                f"{required:.2f}" if required is not None else "—"
            )
            _set_cell_text(row[9 + i], text)


def _add_existing_shielding_table(doc: Any, results: list[PointResult], project: Project) -> bool:
    """Compares existing (as-installed) shielding against what's required.

    Only covers points where ``existing_material``/``existing_thickness`` was
    entered -- returns False (and adds nothing) if none were, so the caller
    can fall back to a placeholder instead of an empty table.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT

    poi_by_id = {poi.id: poi for poi in project.pois}
    rows_data = []
    for result in results:
        poi = poi_by_id.get(result.poi_id)
        if poi is None or not poi.existing_material or poi.existing_thickness <= 0:
            continue
        additional = result.governing_thickness_mm.get(poi.existing_material)
        if additional is None:
            continue
        finding = "Acceptable" if additional <= 0 else f"Add {additional:.2f} mm"
        rows_data.append((result.label, poi.existing_material, poi.existing_thickness, finding))

    if not rows_data:
        return False

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(
        table.rows[0].cells, ["Point", "Material", "As-installed thickness", "Finding"]
    ):
        _set_cell_text(cell, header, bold=True)
    for label, material, thickness, finding in rows_data:
        row = table.add_row().cells
        _set_cell_text(row[0], label)
        _set_cell_text(row[1], material)
        _set_cell_text(row[2], f"{thickness:g}")
        _set_cell_text(row[3], finding)
    return True


def build_report(project: Project, results: list[PointResult]) -> bytes:
    """Build a .docx shielding report from a project's computed results."""
    docx = _docx()
    doc = docx.Document()

    doc.add_heading("Radiation Shielding Report", level=0)

    info_table = doc.add_table(rows=0, cols=2)
    for field, value in [
        ("Facility", None),
        ("Department", None),
        ("Address", None),
        ("Room", project.name),
        ("Unit", None),
        ("Prior Use", None),
        ("Report Date", date.today().isoformat()),
        ("Report Version", "1.0"),
        ("Physicist", None),
        ("Secondary Review", None),
    ]:
        row = info_table.add_row().cells
        _set_cell_text(row[0], field, bold=True)
        if value:
            _set_cell_text(row[1], value)
        else:
            row[1].text = ""
            _placeholder(row[1].paragraphs[0], f"Enter {field.lower()}")

    doc.add_heading("Shielding Overview", level=1)
    for line in [
        "Calculations performed in accordance with industry-standard methodology "
        "described in " + " and ".join(_methods_in_use(project) or ["TG-108 / NCRP 147"]) + ".",
        "Exposure to personnel and general public was taken into consideration.",
        *_goal_lines(project),
        "If any assumptions below are invalid, please notify the reviewing physicist, "
        "as recommendations may differ for different circumstances.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("General Assumptions", level=1)
    for line in _GENERAL_ASSUMPTIONS:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Specific Assumptions", level=1)
    _placeholder_paragraph(
        doc,
        "Enter project-specific assumptions: room construction, workload basis, "
        "prior shielding reports relied upon, and anything else a reviewer would "
        "need to evaluate these results.",
    )

    doc.add_heading("Room Plan and Design Targets", level=1)
    _placeholder_paragraph(
        doc,
        "Insert floor plan screenshot(s) showing source and target locations. "
        "Use the wall color and opacity controls in the app to make barriers "
        "clearly visible before capturing the screenshot.",
    )

    doc.add_heading("Design Targets", level=1)
    if results:
        _add_targets_table(doc, results, project.materials)
    else:
        _placeholder_paragraph(doc, "No points of interest have been evaluated yet.")

    doc.add_heading("Minimum Prescribed Shielding", level=1)
    has_existing = _add_existing_shielding_table(doc, results, project)
    if not has_existing:
        _placeholder_paragraph(
            doc,
            "No existing/as-installed shielding was entered for any point. Enter "
            "existing material and thickness on a point of interest to have this "
            "table compare it against what is required, or fill in this table by hand.",
        )

    doc.add_heading("Summary and Recommendations", level=1)
    needs_shielding = [
        r.label for r in results
        if any(v > 0 for v in r.governing_thickness_mm.values())
    ]
    if needs_shielding:
        doc.add_paragraph(
            "The following points require additional shielding beyond what the "
            "floor plan currently shows: " + ", ".join(needs_shielding) + "."
        )
    else:
        doc.add_paragraph(
            "Based on the barriers currently drawn, no additional shielding is "
            "indicated for the evaluated points -- verify against as-built "
            "conditions before relying on this."
        )
    _placeholder_paragraph(
        doc, "Enter the narrative findings and recommendations for this report."
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    _placeholder(sig, "Signature")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
